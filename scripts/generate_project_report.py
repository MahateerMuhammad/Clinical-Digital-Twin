#!/usr/bin/env python3
"""
scripts/generate_project_report.py
───────────────────────────────────
Generate a Word (.docx) report cataloguing every problem encountered during the
Clinical Digital Twin project — data pipeline, feature engineering, data
integrity, leakage, modelling, environment, and version control — together with
each root cause and the solution applied.

Usage
-----
    .venv/bin/python scripts/generate_project_report.py
    # → reports/Clinical_Digital_Twin_Issues_and_Solutions.docx
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

# ── Palette ───────────────────────────────────────────────────────────────────
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x2E, 0x74, 0xB5)
RED = RGBColor(0xB0, 0x2A, 0x2A)
GREEN = RGBColor(0x1E, 0x7A, 0x3C)
GREY = RGBColor(0x55, 0x55, 0x55)

# ── Content: every problem, grouped by category ───────────────────────────────
# Each item: (title, severity, problem, root_cause, solution)
CATEGORIES = [
    ("1. Data Pipeline & Loading", [
        (
            "Version control was hiding the cleaning pipeline and every model-training module",
            "High",
            "`src/data/` and `src/models/` — 12 files including the entire cleaning pipeline and all "
            "five model trainers — were invisible to git. A fresh clone would have been missing them.",
            "Unanchored `data/` and `models/` entries in `.gitignore` match *any* directory of that "
            "name at any depth, so the rules intended for the top-level dataset and artifact "
            "directories also swallowed the source packages inside `src/`.",
            "Anchored both rules to `/data/` and `/models/` so only the top-level directories are "
            "ignored, and added a test that fails if any tracked source package becomes ignored again. "
            "The failure mode is silent by construction — ignored files simply never appear in "
            "`git status` — so the guard matters more than the fix.",
        ),
        (
            "Chunked CSV reader silently dropped all rows",
            "High",
            "Running the loader with `max_chunks=1` returned an empty table (0 rows).",
            "An off-by-one in `read_csv_chunked`: the loop checked the chunk limit *before* appending, "
            "so a limit of 1 broke out before storing any data.",
            "Reordered to append-then-check, so N chunks yield N chunks of data.",
        ),
        (
            "Load + clean stage took ~17 hours",
            "High",
            "A full load+clean run took 17.3 hours, dominated by tables never used downstream.",
            "Four large tables (`emar`, `emar_detail`, `pharmacy`, `radiology_detail`) — 100M+ rows — "
            "were loaded and cleaned despite zero references in any feature. Per-note `textstat` "
            "readability also ran single-threaded over hundreds of thousands of discharge notes.",
            "Added `pipeline.skip_tables` config to skip the unused tables and a "
            "`compute_note_readability` flag (default off). Verified outputs are byte-identical "
            "since no feature derived from the skipped tables.",
        ),
        (
            "Slow repeated raw CSV re-parsing",
            "Medium",
            "Every pipeline re-run re-parsed multi-GB raw CSVs from scratch.",
            "No caching layer between raw CSV and the cleaning stage.",
            "Added a raw→parquet cache keyed by source mtime + column filter; re-reads skip CSV parsing "
            "while remaining data-identical.",
        ),
    ]),
    ("2. Feature Engineering Correctness", [
        (
            "Valid lab statistic columns were being dropped",
            "High",
            "Lab features for analytes that were present still lost their summary-stat columns.",
            "A `dropna(how='all')` over the whole frame removed columns that were legitimately all-NaN "
            "for absent analytes, taking valid columns with them.",
            "Narrowed the drop to only the all-NaN `slope`/`change` columns; pre-seeded all key labs "
            "with explicit defaults (count=0, missing_ratio=1.0).",
        ),
        (
            "Vitals rolling / trend values mismatched the base statistics",
            "High",
            "For some stays, a vital's rolling-window and trend features came from a different measurement "
            "source than its base mean/min/max.",
            "Base stats used the last itemid present, while rolling/trend used the last itemid with "
            "count ≥ 3 — silently mixing two sources for the same vital.",
            "Split the aggregation: compute base stats and rolling/trend separately, each with its own "
            "`drop_duplicates`, then merge on (stay_id, vital).",
        ),
        (
            "Charlson Comorbidity Index double-counted hierarchical conditions",
            "Medium",
            "The comorbidity index reached an impossible max of 21.",
            "Hierarchical condition pairs (e.g. mild + severe liver disease) were both counted instead of "
            "keeping only the more severe.",
            "Added CCI hierarchy logic to zero the lesser condition when the greater is present; max "
            "dropped to the correct 18.",
        ),
        (
            "Python group-loops made feature building slow",
            "Medium",
            "Lab, vitals, and time-series features used per-group Python loops over tens of millions of rows.",
            "Row-wise / group-wise Python iteration instead of vectorised pandas.",
            "Rewrote to vectorised `groupby.agg` with closed-form slope/least-squares; verified "
            "byte-identical output against captured baselines.",
        ),
    ]),
    ("3. Data Integrity / Corruption", [
        (
            "Test runs overwrote real cleaned data with synthetic fixtures",
            "Critical",
            "Real cleaned tables were destroyed — `chartevents_clean` shrank to 8 rows, `admissions_clean` "
            "to 29 rows — after running the test suite.",
            "`DataCleaner.clean_table` defaulted to `save=True`, and several tests called it without "
            "`save=False`, so synthetic fixtures were written over the real interim parquet files.",
            "Passed `save=False` in all affected tests and recovered the real data from the raw→parquet "
            "cache; regenerated cleaned tables (chartevents 62.8M rows, admissions 546K).",
        ),
        (
            "`hospital_expire_flag` KeyError in the datasets stage",
            "High",
            "The dataset-build step crashed on a missing `hospital_expire_flag` column.",
            "A stale, truncated `admissions_clean.parquet` from an earlier inconsistent state — not a code "
            "bug.",
            "Regenerated `admissions_clean` from cache (full 18 columns); the datasets stage then completed.",
        ),
    ]),
    ("4. Data Leakage (Model Validity)", [
        (
            "Post-hoc ICD / Charlson codes leaking the outcome",
            "Critical",
            "Diagnosis-derived features could encode the label itself (a diagnosis assigned at discharge).",
            "MIMIC-IV ICD codes are finalised at/after discharge; using them to predict in-stay outcomes "
            "leaks the answer, inflating AUROC toward 0.97+.",
            "Established the Run A/B/C leakage protocol; Run B excludes `cci_*`, `dx_*`, and other post-hoc "
            "diagnosis features.",
        ),
        (
            "Discharge-note count features dominated mortality & ICU models",
            "Critical",
            "`sentence_count`, `negation_count`, `medical_keyword_count`, `note_count` were the #1 features "
            "in every model, inflating scores.",
            "These are derived from discharge *summaries* — written at discharge, describing the whole stay "
            "including the outcome. No exclusion list covered them.",
            "Added a shared `NOTE_LEAKAGE` constant wired into the mortality and ICU exclusion lists "
            "(deliberately kept for readmission, where notes precede the 30-day window). Mortality AUROC "
            "0.985→0.980, ICU 0.869→0.847 — leakage removed.",
        ),
        (
            "Full-admission aggregates leaked into in-hospital mortality",
            "High",
            "Whole-stay totals (`medication_count`, `med_duration_hours_*`, `lab_*_last`) are only known at "
            "discharge, correlating with length of stay / the outcome.",
            "Mortality trained on Run B, which retains these full-stay aggregates.",
            "Trained mortality on Run C (`MORTALITY_EXCLUDE_RUN_C`), which strips full-stay aggregates. "
            "AUROC settled at a defensible 0.945.",
        ),
    ]),
    ("5. Modelling & Training", [
        (
            "Hyperparameters were selected on accuracy, not AUROC",
            "Critical",
            "Every model's GroupKFold 'tuning' compared configs by `model.score()` (accuracy).",
            "On imbalanced clinical targets (2–20% positive), accuracy barely moves and rewards predicting "
            "the majority class, so the search was effectively inert.",
            "Built a shared `src/models/tuning.py` that runs Optuna (TPE) maximising patient-grouped "
            "ROC-AUC over a wide search space; wired into all 5 models.",
        ),
        (
            "Narrow 3-configuration grid search",
            "Medium",
            "The 'search' only tried 3 hand-picked hyperparameter sets.",
            "A tiny static grid cannot explore depth / learning-rate / regularisation trade-offs.",
            "Replaced with Optuna (configurable trials via `CDT_TUNE_TRIALS`) over 7–8 hyperparameters.",
        ),
        (
            "No ensemble of the base models",
            "Low",
            "Only individual LR / XGB / LGBM models were reported.",
            "No mechanism combined their complementary strengths.",
            "Added a logistic stacking meta-learner (`blend_predictions`) fit on validation probabilities; "
            "the Stacked Ensemble delivers strong AUPRC with excellent calibration.",
        ),
        (
            "`train_models.py` failed on report export",
            "Medium",
            "`AttributeError: 'str' object has no attribute 'parent'` when writing the comparison report.",
            "`export_model_comparison_markdown` treated `output_path` as a `Path` but received a string.",
            "Coerced with `Path(output_path)` inside the function so any caller works.",
        ),
        (
            "Driver broke on differing model-pipeline APIs",
            "Medium",
            "`too many values to unpack` and `unexpected keyword 'run_type'` when training readmission / ICU.",
            "The pipelines diverge: `prepare_datasets` returns 10 or 11 values; ICU takes no `run_type`; "
            "readmission's `train_*` returns an extra element and its `save_models` differs.",
            "Made the driver introspect signatures, slice `prepared[:10]`, and normalise return arities.",
        ),
        (
            "Report template had hardcoded title and winning model",
            "Low",
            "Every report was titled 'In-Hospital Mortality' and named XGBoost the winner, even for ICU / "
            "readmission where LightGBM won.",
            "The title string and `winning_model_name` were hardcoded.",
            "Report now takes a per-task title and computes the winning model from AUPRC.",
        ),
    ]),
    ("6. Environment & Tooling", [
        (
            "Core scientific libraries missing",
            "Medium",
            "`pyarrow`, `plotly`, and others were unavailable initially.",
            "Fresh environment; a first install attempt timed out on the network.",
            "Installed the full scientific stack into the project virtual environment.",
        ),
        (
            "`shap` could not be installed",
            "Medium",
            "`pip install shap` failed while building `numba`.",
            "`shap` depends on `numba`, which has no wheel for Python 3.14 (too new to build against).",
            "Made every `shap` import lazy and guarded the SHAP plotting step so pipelines train / "
            "evaluate / save normally, skipping only the optional explainability plot.",
        ),
        (
            "LOS runner crashed on a missing module",
            "Medium",
            "`ModuleNotFoundError: No module named 'src.visualization.model_plots'`.",
            "The LOS runner imported ROC/PR, calibration, and SHAP plotting helpers that did not exist.",
            "Created `src/visualization/model_plots.py` with the three helpers and a lazy `shap` import.",
        ),
    ]),
    ("7. Identifier Corruption & Recovery", [
        (
            "Float downcasting silently corrupted patient and admission identifiers",
            "Critical",
            "Laboratory coverage differed systematically between even- and odd-numbered admissions — "
            "roughly half the cohort appeared to have had no bloods taken at all.",
            "A memory optimisation downcast every numeric column to `float32`. Identifiers above "
            "2^24 (16.7M) cannot be represented exactly in float32, so `subject_id` and `hadm_id` "
            "were rounded to even values. Distinct patients collided onto the same identifier and "
            "lab results joined to the wrong admissions.",
            "Excluded identifier columns from downcasting and converted them to pandas `Int32`, which "
            "is exact and still halves the memory of `float64`. Added a magnitude guard for any "
            "unnamed column holding integral values above 2^24. Re-read `labevents` from source CSV "
            "(the parquet cache had been written by the same defective code) and rebuilt all "
            "downstream features. `--verify` asserts odd identifiers are present at ~50% and that "
            "lab coverage is within 5 points between even- and odd-ID admissions.",
        ),
        (
            "The patient split was deliberately NOT regenerated after the repair",
            "Low",
            "A rebuild of this size would normally invalidate every published metric.",
            "Not a defect — a decision that had to be made correctly. Regenerating the split would "
            "have made pre- and post-correction numbers incomparable.",
            "`patient_split.parquet` is keyed on `subject_id`, which was never affected, so the split "
            "was held fixed. Before/after metrics are therefore measured on an identical test cohort, "
            "turning the repair into a controlled experiment rather than a reset.",
        ),
        (
            "The LACE readmission baseline scored below chance",
            "Medium",
            "The published LACE clinical baseline reported AUROC 0.4994 — worse than a coin flip — "
            "which was accepted as a finding about LACE.",
            "The same corrupt laboratory join. LACE depends on lab-derived acuity, so a broken join "
            "removed most of its signal. The number described the defect, not the instrument.",
            "Recomputed after the repair. The lesson recorded alongside it: a baseline scoring below "
            "chance is evidence of a broken input, not of a weak baseline.",
        ),
    ]),
    ("8. Availability Leakage & the Phase 5 Rebuild", [
        (
            "Deterioration model learned 'was a vital sign charted', not the patient's physiology",
            "Critical",
            "The clinical deterioration model reported near-perfect discrimination.",
            "`vital_*` features are recorded per ICU stay. For ward patients who never transferred "
            "they were 100.00% missing (0 of 460,786 rows populated); for those who did transfer they "
            "were 99.99% present. The *presence* of a vital sign was the label. No value needed to be "
            "read for the model to be right.",
            "Excluded all ICU `chartevents`-derived vitals and NEWS2 composites, and rebuilt Phase 5 "
            "as a landmark design: fix T=24h, admit only patients still at risk and event-free at T, "
            "predict transfer within a 48h horizon. Honest performance is AUROC 0.7679 / AUPRC 0.1016 "
            "against a 2.19% base rate — 4.62x enrichment.",
        ),
        (
            "Testing-volume features encoded the clinician's concern, not the patient's state",
            "High",
            "Count and missing-ratio features ranked highly in the deterioration model.",
            "How often a patient is tested reflects how worried the treating team already is. That is "
            "the outcome leaking backwards through clinical behaviour rather than physiology.",
            "Added the full-stay counter and missing-ratio families to the strict exclusion lists, and "
            "an explicit assertion that no volume feature appears in the SHAP top 15.",
        ),
        (
            "The deterioration model was served without its calibrator",
            "Critical",
            "A fully normal patient received a ~79% deterioration risk against a 5.95% base rate.",
            "`LiveModelRunner` had no calibrator entry for deterioration, so the raw booster output "
            "was returned directly. The model is class-weight balanced, which distorts raw scores "
            "badly at low base rates.",
            "Registered the isotonic calibrator. Brier score improved from 0.1636 to 0.0454.",
        ),
        (
            "A file named `..._lightgbm_winning.pkl` contained an XGBoost model",
            "Medium",
            "Anyone reading the filename to learn what was being served got the wrong answer — which "
            "is how the missing calibrator went unnoticed.",
            "The promotion script always copied the true per-task winner into a name that asserted "
            "LightGBM, regardless of which algorithm actually won.",
            "Renamed the promoted artifacts to be algorithm-neutral (`phase5_deterioration_winning.pkl`), "
            "with a legacy-name fallback so existing checkouts keep working.",
        ),
    ]),
    ("9. Serving Layer & Published-Figure Drift", [
        (
            "The LLM layer served pre-correction models with no error raised",
            "Critical",
            "After the identifier repair and retrain, the entire serving layer continued to return "
            "predictions from the old, corrupt models.",
            "`models/` (training output) and `models/best_models/` (what the serving layer loads) used "
            "different naming schemes and nothing bridged them — promotion was a manual copy-and-rename "
            "that had not been done.",
            "`scripts/maintenance/promote_models.py` now performs the mapping explicitly and archives "
            "what it replaces.",
        ),
        (
            "Clinicians were told '6-hour deterioration risk' by a model predicting 48 hours",
            "Critical",
            "The clinical report described a 6-hour deterioration horizon while the served model "
            "predicted ICU transfer within 48 hours, for patients already stable for 24 hours. A "
            "reader seeing 3% would have understood something eight times more urgent than the model "
            "actually claimed.",
            "`phase5_deterioration_window_hours: 6.0` was the lead time of the superseded case-control "
            "design. It survived the landmark rebuild untouched because it lived in a constants table "
            "nobody had reason to revisit.",
            "Replaced with `phase5_landmark_hours` and `phase5_horizon_hours` read from the promoted "
            "model's own metadata, and corrected every label in the prompt builder, report composer "
            "and clinical assistant.",
        ),
        (
            "A published notebook made three false claims about the model it described",
            "Critical",
            "`11_deterioration_baseline.ipynb` presented the superseded 6-hour case-control design, "
            "claimed 'AUROC >0.99 and AUPRC >0.99', and attributed that performance to "
            "`vital_heart_rate_mean`, `vital_sbp_max` and NEWS2 composites — features the matrix does "
            "not contain at all.",
            "The notebook *stated* its parameters and metrics as literal text rather than deriving "
            "them, so the Phase 5 rebuild left it describing a model that no longer existed.",
            "Rewrote the generator so the notebook reads landmark, horizon, winner and feature count "
            "from `phase5_deterioration_landmark.json` and recomputes every metric from the served "
            "artifacts. It now reproduces the published table exactly rather than quoting it, and "
            "asserts its own leakage invariants.",
        ),
        (
            "Generated reports carried hand-written narrative that contradicted their own tables",
            "Medium",
            "The payload-fidelity report explained that ICU admission's rank correlation 'is near "
            "zero' while the table beside it showed +0.79.",
            "The prose was a fixed string in the generator, describing the results as they stood on "
            "the day it was written. The table updated; the paragraph explaining it did not.",
            "The narrative is now computed from the measurement — ranges, extremes and counts are "
            "derived, so the explanation cannot disagree with the data it explains.",
        ),
        (
            "A quoted twin-retrieval AUROC was a literal that had gone stale",
            "Low",
            "0.7253 was quoted to clinicians in the twin section of the structured prompt.",
            "The figure was hardcoded in the prompt builder rather than read from its evaluation table.",
            "Moved into the system-constants table with a pointer to the source report, alongside a "
            "test that fails if the two disagree.",
        ),
    ]),
    ("10. Emergency Department Integration", [
        (
            "MIMIC-IV-ED coverage was overestimated by nearly a factor of two",
            "Medium",
            "The expected gain from adding ED data was reported as 'roughly 69.5% of admissions, "
            "lifting vitals coverage to ~75%'. The true ED module coverage is 37.1%.",
            "The 69.5% figure came from `admission_location == 'EMERGENCY ROOM'` in the admissions "
            "table — which records where a patient came from, not whether a linked ED record exists. "
            "Even among emergency-room admissions only 56.5% have one.",
            "Measured the join directly before writing any feature code: 202,415 of 546,028 admissions "
            "link to an ED stay. The gain is 0% → 37%, not 16% → 75%, and is documented as such.",
        ),
        (
            "Impossible vital signs survived the shared physiological ranges",
            "Medium",
            "ED triage records showed diastolic BP of 0.0 and systolic of 1.0 after clipping.",
            "The shared `VITAL_RANGES` lower bounds are inclusive of zero, so documented blood-pressure "
            "cuff-failure artefacts passed the filter unchanged.",
            "Added `ED_VITAL_FLOORS` applied locally to the ED path only, so the ICU path and every "
            "already-trained phase keep their exact fitted thresholds. Verified ranges afterwards: "
            "SBP 24–274, DBP 10–247, HR 14–207.",
        ),
        (
            "41% of serial ED observations were charted at or after admission time",
            "Medium",
            "A naive per-stay window would have admitted post-admission measurements into a "
            "pre-admission feature set.",
            "Hospital registration happens while the patient is still physically in the ED, so "
            "`admittime` falls inside the ED stay and `outtime` follows it 99.7% of the time.",
            "Kept a strict per-`charttime` cut against `admittime` as the default, and added an "
            "explicit `window_hours` parameter so the 24-hour protocol variant is a flag rather than "
            "a rewrite.",
        ),
        (
            "ED availability was checked for leakage rather than assumed safe",
            "Low",
            "ED features are populated for only 37% of the cohort, which is exactly the missingness "
            "pattern that caused the Phase 5 failure.",
            "Nothing was wrong — but the risk had to be measured, not argued.",
            "`ed_available` alone scores AUROC 0.5097 for mortality, 0.5018 for ICU admission and "
            "0.4985 for readmission. The Phase 5 failure mode does not repeat. The guard is retained "
            "as cheap insurance and pinned by test.",
        ),
    ]),
    ("11. Payload-Based Serving (Unseen Patients)", [
        (
            "Four of five models could not be served for an unseen patient",
            "Critical",
            "Only in-hospital mortality could be answered from a presentation payload. Readmission "
            "(32.6% retention), ICU admission (24.0%), length of stay (−2.7%) and deterioration "
            "(−0.8%) were all withheld — two of them ranking patients backwards.",
            "Two causes, both in the serving path rather than in the models. The converter wrote "
            "booster column names by hand instead of emitting source categoricals and letting the "
            "fitted encoder expand them; and the payload schema never asked for race, language, "
            "insurance, marital status, admission type/location or prior utilisation. Those one-hot "
            "families are 86 of the mortality model's 164 features, and `prior_*` is the block the "
            "readmission model was built on — it was being asked its question with the patient's own "
            "readmission history withheld.",
            "Routed payloads through `encode_admission_frame`, the same encoder used at fit time, and "
            "extended the schema as recommended (not required) fields. Coverage rose 18.3% → 67.7%. "
            "Retention: mortality 85.6%, readmission 81.7%, ICU 75.6%, deterioration 91.6% — four of "
            "five tasks now served, with the 66.7% floor left untouched.",
        ),
        (
            "A female patient reached the deterioration model as 'not male, sex unknown'",
            "High",
            "`gender_M` was written directly as 0.0 while `gender_F` — which Phase 5 also expects — "
            "was left absent, i.e. missing rather than zero.",
            "The same hand-written dummy names. One-hot families were emitted with a single member "
            "set and every sibling left unfilled.",
            "Fixed by the encoder change above, and pinned by a test asserting both sides of the "
            "family are populated.",
        ),
        (
            "Absent one-hot columns were treated as unknown when they were known zeros",
            "High",
            "A payload stating `race = WHITE` left the other 31 `race_*` features missing — asserting "
            "the patient's race was unknown immediately after it had been given.",
            "The project's rule that unsupplied features must be NaN rather than 0.0 is correct for "
            "measurements, and was applied uniformly. For one-hot families it inverts: the zero *is* "
            "what the supplied category determined.",
            "Absent dummies of a supplied category now fill 0.0, tracked per column and per row; "
            "absent dummies of a category nobody mentioned stay NaN. Both directions are pinned by "
            "test, including the reference level dropped at fit time.",
        ),
        (
            "The fidelity harness reimplemented the production converter instead of calling it",
            "Medium",
            "The measurement that decides which models may be served kept its own copy of the payload "
            "schema.",
            "A mirror that maintains its own copy stops being a mirror. It would have reported no "
            "improvement from a change that had already landed in production.",
            "The harness now reads the field lists from the runner, so the two cannot diverge.",
        ),
        (
            "The evaluation harness mapped labs onto column names that did not exist",
            "High",
            "Every laboratory value fell through to a clinically-normal constant, making all evaluated "
            "payloads near-identical — so the harness was measuring the constants, not the cohort.",
            "`lab_creatinine_max_24h` and five siblings were assumed to exist; the windowed build emits "
            "first/last rather than max.",
            "Corrected the mapping and recorded that any result predating 2026-08-01 describes the "
            "corpus rather than the patients.",
        ),
        (
            "The coverage guard suppressed every prediction it was meant to protect",
            "Medium",
            "A 30% minimum-coverage floor fired on 100% of payloads, discarding the mortality estimate "
            "— the one output measured to retain most of its validated discrimination — alongside the "
            "four that genuinely did not.",
            "The floor was chosen against the earlier zero-filling defect. Once unsupplied features "
            "became NaN, coverage was no longer the right instrument: it counts how much input is "
            "missing, not whether what remains still discriminates.",
            "Lowered to a 10% backstop and made per-task withholding a function of measured retention. "
            "The floor is explicitly not re-tuned to track coverage, so it stays a backstop rather than "
            "a second competing gate.",
        ),
    ]),
    ("12. Inert Guards & Unit Errors", [
        (
            "Two leakage exclusions matched nothing and protected nothing",
            "High",
            "`ICU_ADMISSION_EXCLUDE` and `LOS_EXCLUDE_STRICT` listed `fluids_*` and `vitals_*`. The "
            "built columns are `fluid_*` and `vital_*` — singular — so neither pattern ever matched.",
            "A plural/singular slip. An exclusion that matches nothing is indistinguishable from one "
            "that matches everything: both are silent.",
            "Corrected to the singular forms. Nothing leaked, because the training matrix happens to "
            "contain no such column — but had vitals or fluids ever been adopted, Phase 3 would have "
            "predicted ICU admission from ICU-only features with the guard silent. Added a test that "
            "checks every pattern in every exclusion list against the full namespace of columns the "
            "project can build (434 subtests); a pattern matching nothing now fails, and exemptions "
            "must carry a written reason.",
        ),
        (
            "Fluid totals added millilitres to milligrams",
            "Medium",
            "`fluid_input_total` summed the raw amount column across 22 different units. Only 54% of "
            "rows are millilitres; the rest are mg, mcg, grams, doses, units and mEq. Since mcg and "
            "grams differ by a factor of a million, one microgram-dosed infusion could dominate a "
            "patient's apparent fluid intake. The median intake was inflated by ~47%.",
            "The builder never read `amountuom`.",
            "Replaced with a unit→millilitre conversion table. Rebuilt the artifact: median input "
            "8,669 → 5,886 mL, median balance 4,086 → 1,360. No model retrained — fluid features "
            "reach no model, which is the only reason this was harmless.",
        ),
        (
            "The first attempt at the fluid fix was itself wrong, twice",
            "Low",
            "An allow-list accepted litres and summed them as millilitres, understating those rows "
            "a thousandfold. After that was fixed, the micro-litre key was unreachable.",
            "The first error was treating 'is a volume' as sufficient when 'is the same volume unit' "
            "was required. The second: Python's `casefold()` maps the MICRO SIGN (U+00B5, which is "
            "what MIMIC stores) to GREEK SMALL LETTER MU, so a literal `µl` key never matches a "
            "casefolded lookup — those rows would have been dropped as unrecognised while looking "
            "perfectly correct in the source.",
            "Converted rather than filtered, and normalised the lookup table through the same "
            "`casefold()` the data passes through. Both traps were caught by tests written before the "
            "artifact was rebuilt, which is the reason they are recorded here rather than shipped.",
        ),
    ]),
    ("13. Documentation & Provenance Drift", [
        (
            "A notebook published metrics computed on a different machine, before a data repair",
            "High",
            "`10_los_two_stage.ipynb` carried stored outputs reporting AUROC 0.8114 while the "
            "corresponding phase report showed 0.8997.",
            "The notebook was executed before the laboratory join was repaired and never re-run. Its "
            "saved cell outputs are a snapshot of a superseded state that looks authoritative.",
            "Added provenance banners recording when each notebook was executed, against which data, "
            "and which figures supersede it — plus a test that fails if a notebook's stored metrics "
            "drift from the promoted models.",
        ),
        (
            "Emergency-department features were built, tested, and used by nothing",
            "Medium",
            "66 ED features passed every leakage screen and reached no model. An artifact that exists "
            "but is unused reads exactly like one in production.",
            "ED integration completed after the last full pipeline run, and adopting it requires "
            "rebuilding the modelling matrix and retraining Phases 1–5.",
            "Made the status explicit and self-updating: the coverage report derives 'staged' versus "
            "'in use' by counting `ed_` columns in the selected matrix, so it flips automatically on "
            "adoption rather than relying on someone remembering to edit it.",
        ),
        (
            "No phase report disclosed that the feature matrix contains no vital signs",
            "Medium",
            "A reader of the Phase 1–4 reports would reasonably assume vital signs were represented. "
            "The matrix contains zero `vital_*` columns.",
            "The exclusion was made for a good reason (ICU-only availability leakage) and recorded in "
            "Phase 5, but never surfaced in the phases that also inherit it.",
            "Added the disclosure to the affected reports, derived from the matrix rather than "
            "asserted, with a test that fails if the two disagree.",
        ),
        (
            "Feature engineering choices carried no literature justification",
            "Medium",
            "The feature engineering report contained zero citations for derived features making "
            "physiological claims — shock index, pulse pressure, NEWS2 components, the 24-hour window, "
            "the physiological clipping ranges.",
            "The report documented *what* was computed but not *why* those choices are clinically "
            "defensible.",
            "Added a clinical rationale and literature section covering the derived features and the "
            "windowing decision.",
        ),
    ]),
    ("14. Version Control", [
        (
            "Model pickles were committed despite a .gitignore rule",
            "Medium",
            "12 `.pkl` model files were tracked and pushed to GitHub.",
            "The ignore rule was `.pkl` (matches only a file literally named '.pkl'), not `*.pkl`, so "
            "`models/*.pkl` was never ignored.",
            "Corrected to `*.pkl` plus `models/`, and untracked the files with `git rm --cached` (kept on "
            "disk).",
        ),
        (
            "Repository needed to move to a new owner",
            "Low",
            "The project had to be re-hosted under a different GitHub account, keeping history.",
            "Original remote pointed at a different owner's repository.",
            "Installed GitHub CLI, created a new private repo under the target account, swapped the remote, "
            "and pushed with full commit history preserved.",
        ),
    ]),
]


# ── Rendering helpers ─────────────────────────────────────────────────────────
def sev_color(sev: str) -> RGBColor:
    return {"Critical": RED, "High": RGBColor(0xC0, 0x60, 0x10),
            "Medium": ACCENT, "Low": GREEN}.get(sev, GREY)


def add_heading_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("—" * 46)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)


def build_document() -> Document:
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ── Title page ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title.add_run("Clinical Digital Twin")
    t.bold = True
    t.font.size = Pt(26)
    t.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Problems Encountered & Solutions Applied")
    s.font.size = Pt(15)
    s.font.color.rgb = ACCENT

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = meta.add_run(f"MIMIC-IV Preprocessing & Clinical Prediction Pipeline\nGenerated {date.today().isoformat()}")
    m.font.size = Pt(10.5)
    m.font.color.rgb = GREY

    total = sum(len(items) for _, items in CATEGORIES)
    crit = sum(1 for _, items in CATEGORIES for it in items if it[1] == "Critical")

    # Executive summary
    doc.add_paragraph()
    h = doc.add_heading("Executive Summary", level=1)
    h.runs[0].font.color.rgb = NAVY
    doc.add_paragraph(
        f"This report documents {total} distinct problems resolved across the lifecycle of the "
        f"Clinical Digital Twin project, spanning the data pipeline, feature engineering, data integrity, "
        f"model-validity (leakage), model training, environment/tooling, and version control. "
        f"Of these, {crit} were critical — most notably several forms of target leakage that had inflated "
        f"model performance, and a data-corruption incident that was fully recovered from cache. "
        f"The corrected models report honest, leakage-audited, literature-consistent performance."
    )

    # Category counts table
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    hdr[0].paragraphs[0].add_run("Category").bold = True
    hdr[1].paragraphs[0].add_run("Issues").bold = True
    for name, items in CATEGORIES:
        row = tbl.add_row().cells
        row[0].text = name
        row[1].text = str(len(items))

    doc.add_page_break()

    # ── Detailed sections ──
    for name, items in CATEGORIES:
        h = doc.add_heading(name, level=1)
        h.runs[0].font.color.rgb = NAVY
        for (ptitle, sev, problem, cause, solution) in items:
            hp = doc.add_heading(ptitle, level=2)
            hp.runs[0].font.color.rgb = ACCENT
            hp.runs[0].font.size = Pt(12.5)

            sp = doc.add_paragraph()
            lbl = sp.add_run("Severity: ")
            lbl.bold = True
            sv = sp.add_run(sev)
            sv.bold = True
            sv.font.color.rgb = sev_color(sev)

            for field, text in (("Problem", problem), ("Root Cause", cause), ("Solution", solution)):
                para = doc.add_paragraph()
                r = para.add_run(f"{field}:  ")
                r.bold = True
                r.font.color.rgb = GREEN if field == "Solution" else NAVY
                para.add_run(text)
            add_heading_rule(doc)
        doc.add_page_break()

    # ── Closing note ──
    h = doc.add_heading("Outcome", level=1)
    h.runs[0].font.color.rgb = NAVY
    doc.add_paragraph(
        "After these fixes, the pipeline runs end-to-end reproducibly, the interim data is protected "
        "against accidental overwrite, identifiers survive round-tripping exactly, and the five "
        "clinical models (in-hospital mortality, 30-day readmission, ICU admission, length of stay, "
        "and clinical deterioration) are trained with AUROC-based Optuna tuning, patient-grouped "
        "cross-validation, isotonic calibration, and a stacking ensemble. Most importantly, the "
        "removal of note-based, full-stay-aggregate and availability leakage replaced inflated scores "
        "with honest, defensible, literature-consistent performance:"
    )
    for line in (
        "In-Hospital Mortality — AUROC 0.945 (leakage-clean Run C)",
        "ICU Admission — AUROC 0.921",
        "Hospital Length of Stay — AUROC 0.900 (Stage A)",
        "30-Day Readmission — AUROC 0.706 (at the genuine feature ceiling)",
        "Clinical Deterioration — AUROC 0.768, AUPRC 0.102 at a 2.19% base rate (4.62x enrichment), "
        "rebuilt as a 24-hour landmark with a 48-hour horizon",
    ):
        b = doc.add_paragraph(style="List Bullet")
        b.add_run(line)

    doc.add_paragraph()
    doc.add_paragraph(
        "Four of the five models can now be served for a patient who is not in the cohort, from a "
        "presentation payload alone — up from one — with length of stay withheld at 57.9% retention "
        "because discharge planning and social circumstances are not observable at presentation."
    )

    doc.add_paragraph()
    h = doc.add_heading("The Recurring Defect", level=1)
    h.runs[0].font.color.rgb = NAVY
    doc.add_paragraph(
        "Most of the critical entries above are the same defect wearing different clothes: a value "
        "living somewhere that does not update when its source changes. A horizon in a constants "
        "table, a metric typed into a notebook, an AUROC quoted in a prompt, a paragraph explaining a "
        "table it no longer matches, an exclusion pattern spelled plural against singular columns. "
        "Each was individually small and none produced an error message."
    )
    doc.add_paragraph(
        "The durable fix in every case was the same, and it was not vigilance: make the artifact "
        "derive the value rather than restate it, and add a test that fails when the derivation and "
        "the source disagree. Reports now compute their own narrative, notebooks read their "
        "parameters from the promoted model metadata, status lines are counted rather than asserted, "
        "and exclusion patterns are validated against the real column namespace. A guard that can "
        "match nothing is now a test failure rather than a silence."
    )

    return doc


def main() -> None:
    out_dir = Path("reports")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Clinical_Digital_Twin_Issues_and_Solutions.docx"
    doc = build_document()
    doc.save(out_path)
    n = sum(len(items) for _, items in CATEGORIES)
    print(f"✓ Word document generated: {out_path}")
    print(f"  {len(CATEGORIES)} categories, {n} documented problems + solutions.")


if __name__ == "__main__":
    main()
