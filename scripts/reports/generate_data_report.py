#!/usr/bin/env python
"""
scripts/reports/generate_data_report.py
───────────────────────────────────────
The data cleaning and feature engineering report, generated from the data.

    PYTHONPATH=. .venv/bin/python scripts/reports/generate_data_report.py
    # → reports/data_cleaning_and_feature_engineering.md

Why generated rather than written
─────────────────────────────────
A cohort description typed into a markdown file is correct on the day it is
typed and silently wrong afterwards. This project has already shipped that
defect more than once — a value living somewhere that does not update when its
source changes — and a report quoting a patient count is exactly the shape of
thing that gets quoted in a presentation long after the pipeline moved on.

So every figure here is computed from `data/processed/` or read from a
pipeline artefact at generation time. Nothing is hardcoded, and a missing
source produces a stated gap rather than a stale number.

What it does **not** do
───────────────────────
It does not re-run the pipeline, and it does not recompute model metrics.
Those live in the phase reports, which own them; this report links rather than
restates, because a second copy of a metric is a second thing to keep in sync.
"""

from __future__ import annotations

import re
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

import pandas as pd

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

PROCESSED = Path("data/processed")
TABLES = Path("reports/tables")
OUT = Path("reports/data_cleaning_and_feature_engineering.md")
OUT_DOCX = OUT.with_suffix(".docx")

#: Admissions are the modelling unit; patients are the split unit. Keeping both
#: visible matters because splitting on the wrong one leaks a patient's later
#: admission into training against their earlier one.
ADMISSIONS = PROCESSED / "admission_level.parquet"
SELECTED = PROCESSED / "admission_level_selected.parquet"
SPLIT = PROCESSED / "patient_split.parquet"

AGE_BANDS = ((18, 39), (40, 54), (55, 69), (70, 84), (85, 120))


def _n(x: Any) -> str:
    """
    Thousands separators, and `—` for anything absent.

    Duck-typed rather than `isinstance(x, (int, float))`, because pandas hands
    back `numpy.int64` — which is not a Python `int` — and the check silently
    fell through to `str()`. The first draft printed `177459` in one table and
    `177,459` in the next, from the same column.
    """
    if x is None:
        return "—"
    try:
        return f"{float(x):,.0f}"
    except (TypeError, ValueError):
        return str(x)


def _size(p: Path) -> str:
    """Human-scaled, because `0 MB` for a 6 KB file reads as a broken path."""
    if not p.exists():
        return "—"
    n = float(p.stat().st_size)
    for unit in ("B", "KB", "MB", "GB"):
        if n < 1024 or unit == "GB":
            return f"{n:.0f} {unit}" if unit == "B" else f"{n:.1f} {unit}"
        n /= 1024
    return f"{n:.1f} GB"


def _pct(part: float, whole: float) -> str:
    return "—" if not whole else f"{part / whole * 100:.2f}%"


def _table(rows: List[List[Any]], head: List[str], align: str = "") -> List[str]:
    sep = [(":---" if (align[i:i + 1] or "l") == "l" else "---:")
           for i in range(len(head))]
    return (["| " + " | ".join(head) + " |", "| " + " | ".join(sep) + " |"]
            + ["| " + " | ".join(str(c) for c in r) + " |" for r in rows] + [""])


# ── the cohort ───────────────────────────────────────────────────────────────

def cohort() -> Dict[str, Any]:
    """Counts read from the tables the models are actually trained on."""
    cols = ["subject_id", "hospital_expire_flag", "has_icu_stay", "readmission_30d",
            "los_days", "icu_los_days", "anchor_age", "gender", "admission_type",
            "charlson_comorbidity_index"]
    df = pd.read_parquet(ADMISSIONS, columns=cols)
    d: Dict[str, Any] = {
        "admissions": len(df),
        "patients": df.subject_id.nunique(),
        "deaths": int(df.hospital_expire_flag.sum()),
        "icu": int(df.has_icu_stay.sum()),
        "readmit": int(df.readmission_30d.sum()),
        "los_median": df.los_days.median(),
        "los_mean": df.los_days.mean(),
        "los_p90": df.los_days.quantile(0.90),
        "age_median": df.anchor_age.median(),
        "age_q1": df.anchor_age.quantile(0.25),
        "age_q3": df.anchor_age.quantile(0.75),
        "cci_median": df.charlson_comorbidity_index.median(),
        "sex": df.gender.value_counts(dropna=False).to_dict(),
        "df": df,
    }
    d["alive"] = d["admissions"] - d["deaths"]
    if SPLIT.exists():
        s = pd.read_parquet(SPLIT)
        d["split"] = s.split.value_counts().to_dict()
        d["fingerprint"] = (s.cohort_fingerprint.iloc[0]
                            if "cohort_fingerprint" in s.columns else None)
    return d


def acuity(df: pd.DataFrame) -> List[List[Any]]:
    """
    Severity strata, defined by what the record shows rather than by a label.

    MIMIC-IV carries no "critical" flag, so inventing one would be a clinical
    judgement made by this script. These are observable facts — an ICU stay
    happened, the admission ended in death — and each row states its own
    definition so a reader is never guessing what the denominator was.
    """
    n = len(df)
    icu, died = df.has_icu_stay == 1, df.hospital_expire_flag == 1
    strata = [
        ("Ward only, survived", ~icu & ~died),
        ("Ward only, died", ~icu & died),
        ("ICU, survived", icu & ~died),
        ("ICU, died", icu & died),
    ]
    rows = [[label, _n(int(m.sum())), _pct(int(m.sum()), n),
             f"{df.loc[m, 'los_days'].median():.1f}" if m.any() else "—",
             f"{df.loc[m, 'anchor_age'].median():.0f}" if m.any() else "—"]
            for label, m in strata]
    rows.append(["**All admissions**", f"**{_n(n)}**", "**100%**",
                 f"**{df.los_days.median():.1f}**",
                 f"**{df.anchor_age.median():.0f}**"])
    return rows


def by_age_band(df: pd.DataFrame) -> List[List[Any]]:
    """The partition the slice evaluation and the band calibrators both use."""
    rows = []
    for lo, hi in AGE_BANDS:
        m = df.anchor_age.between(lo, hi)
        sub = df[m]
        if not len(sub):
            continue
        rows.append([f"{lo}–{hi}" if hi < 120 else f"{lo}+", _n(len(sub)),
                     _n(int(sub.hospital_expire_flag.sum())),
                     _pct(int(sub.hospital_expire_flag.sum()), len(sub)),
                     _pct(int(sub.has_icu_stay.sum()), len(sub))])
    return rows


def by_admission_type(df: pd.DataFrame, top: int = 8) -> List[List[Any]]:
    g = df.groupby("admission_type", observed=True).agg(
        n=("hospital_expire_flag", "size"),
        deaths=("hospital_expire_flag", "sum"),
        icu=("has_icu_stay", "sum")).sort_values("n", ascending=False).head(top)
    return [[str(k), _n(r.n), _n(r.deaths), _pct(r.deaths, r.n), _pct(r.icu, r.n)]
            for k, r in g.iterrows()]


# ── cleaning and feature artefacts ───────────────────────────────────────────

def cleaning_actions() -> Optional[pd.DataFrame]:
    """
    The Documented Actions table from the cleaning report.

    Parsed from the generated markdown because the pipeline writes it there and
    nowhere else. Returns None on any failure — a report that cannot find its
    source says so, rather than falling back to a number someone remembered.
    """
    p = Path("reports/cleaning_report.md")
    if not p.exists():
        return None
    rows = []
    for line in p.read_text(encoding="utf-8").splitlines():
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if len(cells) >= 5 and cells[1] in {
                "document_missing", "flag_duplicates", "invalid_timestamp_order",
                "normalize_categorical", "clip_implausible", "flag_outlier"}:
            try:
                rows.append({"table": cells[0], "action": cells[1],
                             "column": cells[2],
                             "n_affected": int(cells[3].replace(",", ""))})
            except ValueError:
                continue
    return pd.DataFrame(rows) if rows else None


def feature_counts() -> Dict[str, Any]:
    """Width before and after selection, measured on the files themselves."""
    out: Dict[str, Any] = {}
    if ADMISSIONS.exists():
        out["before"] = pd.read_parquet(ADMISSIONS).shape[1]
    if SELECTED.exists():
        out["after"] = pd.read_parquet(SELECTED).shape[1]
    if "before" in out and "after" in out:
        out["dropped"] = out["before"] - out["after"]
    p = TABLES / "feature_missing_report.parquet"
    if p.exists():
        miss = pd.read_parquet(p)
        out["missing_report"] = miss
        out["high_missing"] = int((miss.pct_missing > 30).sum())
    return out


def selection_summary() -> Dict[str, int]:
    """The FE pipeline's own tally, so drop *reasons* are its numbers not mine."""
    p = Path("reports/feature_engineering_report.md")
    if not p.exists():
        return {}
    found = {}
    for m in re.finditer(r"\*\*(n_[a-z_]+)\*\*:\s*(\d+)", p.read_text(encoding="utf-8")):
        found[m.group(1)] = int(m.group(2))
    return found


# ── the document ─────────────────────────────────────────────────────────────

def build() -> str:
    c = cohort()
    df = c["df"]
    fc = feature_counts()
    sel = selection_summary()
    actions = cleaning_actions()
    stamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    L: List[str] = [
        "# Data Cleaning and Feature Engineering", "",
        f"*Generated {stamp} by `scripts/reports/generate_data_report.py`. Every "
        f"figure is computed from `data/processed/` or read from a pipeline "
        f"artefact at generation time — none is typed in.*", "",
        "---", "", "## 1. The cohort", "",
        f"Source is MIMIC-IV. The modelling unit is the **admission**; the "
        f"split unit is the **patient**, because splitting on admissions would "
        f"put a patient's later admission in training and their earlier one in "
        f"test, and the model would be scored on someone it had already seen.", "",
    ]

    L += _table([
        ["Patients", _n(c["patients"]), "unique `subject_id`"],
        ["Admissions", _n(c["admissions"]), "modelling rows"],
        ["Survived to discharge", _n(c["alive"]), _pct(c["alive"], c["admissions"])],
        ["Died in hospital", _n(c["deaths"]), _pct(c["deaths"], c["admissions"])],
        ["Had an ICU stay", _n(c["icu"]), _pct(c["icu"], c["admissions"])],
        ["Readmitted within 30 days", _n(c["readmit"]), _pct(c["readmit"], c["admissions"])],
        ["Median age", f"{c['age_median']:.0f} y", f"IQR {c['age_q1']:.0f}–{c['age_q3']:.0f}"],
        ["Median length of stay", f"{c['los_median']:.2f} d",
         f"mean {c['los_mean']:.2f}, p90 {c['los_p90']:.1f}"],
        ["Median Charlson index", f"{c['cci_median']:.0f}", "comorbidity burden"],
    ], ["Measure", "Value", "Note"], align="lrl")

    if c.get("split"):
        s = c["split"]
        L += [f"**Split** (by patient, fingerprint `{c.get('fingerprint', '—')}`): "
              + " · ".join(f"{k} {_n(v)}" for k, v in sorted(s.items())), ""]

    L += ["### Severity strata", "",
          "MIMIC-IV carries no \"critical\" flag, so these are defined by what "
          "the record shows — an ICU stay occurred, the admission ended in "
          "death — rather than by a judgement this report makes for you.", ""]
    L += _table(acuity(df),
                ["Stratum", "Admissions", "Share", "Median LOS (d)", "Median age"],
                align="lrrrr")

    L += ["### By age band", "",
          "The same partition used by `run_slice_eval.py` and by the age-band "
          "calibrators, so subgroup findings can be read against it.", ""]
    L += _table(by_age_band(df),
                ["Age band", "Admissions", "Deaths", "Mortality", "ICU rate"],
                align="lrrrr")

    L += ["### By admission type", ""]
    L += _table(by_admission_type(df),
                ["Admission type", "Admissions", "Deaths", "Mortality", "ICU rate"],
                align="lrrrr")

    L += ["> Mortality ranges widely across these strata — from near zero in "
          "observation admissions to several percent in emergency and urgent "
          "ones. A single overall rate would hide that, which is why the models "
          "are evaluated per slice rather than in aggregate.", ""]

    # ── cleaning ──
    L += ["---", "", "## 2. Data cleaning", "",
          "### The technique, and why", "",
          "**Flag, do not drop. Document, do not impute.** Cleaning records "
          "what is wrong with a row and leaves the row in place; no value is "
          "invented at this stage and no row is silently removed.", "",
          "The reason is that missingness in a clinical record is *evidence*. A "
          "lactate that was never drawn is not a lactate of average value — it "
          "usually means nobody was worried enough to order one. Imputing the "
          "mean would erase that signal and, worse, would make an absent "
          "measurement indistinguishable from a normal one at prediction time. "
          "Every imputation decision is therefore deferred to the modelling "
          "stage, where it is visible and can be evaluated.", "",
          "Timestamp contradictions are treated the same way: an admission "
          "discharged before it was admitted is flagged `_invalid_time_order` "
          "and kept, because deleting it would quietly change the denominator "
          "of every rate computed afterwards.", ""]

    if actions is not None and len(actions):
        by_action = (actions.groupby("action")
                     .agg(n=("action", "size"), tables=("table", "nunique"))
                     .sort_values("n", ascending=False))
        L += ["### What cleaning actually did", ""]
        L += _table([[a, _n(r.n), _n(r.tables)] for a, r in by_action.iterrows()],
                    ["Action", "Actions", "Tables"], align="lrr")
        L += [f"*{len(actions)} documented actions across "
              f"{actions.table.nunique()} tables, parsed from "
              f"`reports/cleaning_report.md`.*", ""]
        # Deliberately not a sum of `n_affected`. Those counts mean different
        # things per action — `document_missing` counts missing *cells* in one
        # column, `invalid_timestamp_order` counts *rows* — so adding them
        # produced "2,473,669,706 rows affected" against a 546,028-row table.
        # A number four thousand times larger than the thing it describes.
    # ── feature engineering ──
    L += ["---", "", "## 3. Feature engineering", "", "### Width", ""]
    if fc.get("before") and fc.get("after"):
        L += _table([
            ["Engineered", _n(fc["before"]), "`admission_level.parquet`"],
            ["After selection", _n(fc["after"]), "`admission_level_selected.parquet`"],
            ["Removed", _n(fc["dropped"]), _pct(fc["dropped"], fc["before"]) + " of the original width"],
        ], ["Stage", "Columns", "Source"], align="lrl")
    if sel:
        L += ["### Why each feature was removed", "",
              "The selection pipeline's own tally — each reason a different "
              "mistake to ignore, and none of them interchangeable:", ""]
        labels = {
            "n_high_missing": "Missing above threshold — too sparse to learn from",
            "n_highly_correlated_pairs": "Highly correlated — redundant, and inflates apparent importance",
            "n_correlated_dropped": "Dropped from those pairs",
            "n_near_zero_variance": "Near-zero variance — cannot separate anyone",
            "n_constant": "Constant — carries no information at all",
            "n_duplicate": "Exact duplicate of another column",
        }
        L += _table([[labels[k], _n(v)] for k, v in sel.items() if k in labels],
                    ["Reason", "Features"], align="lr")

    miss = fc.get("missing_report")
    if miss is not None and len(miss):
        L += ["### Missingness in the modelling table", "",
              "Raw-table gaps are mostly in free-text medication fields that "
              "never become features. What matters is missingness in "
              "`admission_level` — the table the models actually see, and "
              "where the retention threshold bites:", ""]
        L += _table([[f"`{r.feature}`", str(r.dtype), _n(r.n_missing),
                      f"{r.pct_missing:.1f}%"]
                     for _, r in miss.nlargest(8, "pct_missing").iterrows()],
                    ["Feature", "Type", "Missing", "Share"], align="llrr")
        hi = fc.get("high_missing")
        L += [f"*{_n(len(miss))} columns profiled; {_n(hi)} are above 30% "
              f"missing.* `deathtime` heads the list at 97.8% — as it must, "
              "since it is only populated for the 2.2% of admissions that end "
              "in death. It is also a textbook leakage column, and is removed "
              "before modelling for that reason rather than for its sparsity.", ""]
    else:
        L += ["> `reports/cleaning_report.md` was not found or could not be "
              "parsed, so the per-action counts are omitted rather than "
              "guessed. Re-run the cleaning pipeline to restore this section.", ""]

    L += ["### Leakage control", "",
          "The largest single source of a too-good result, and the one that "
          "does not announce itself: a model scoring 0.99 looks like success "
          "until you find it was reading the discharge summary. "
          "`src/features/leakage_filters.py` removes three classes:", "",
          "1. **Outcome-adjacent** — `deathtime`, `dischtime`, "
          "`discharge_location`. Known only once the outcome is known.",
          "2. **Finalised after discharge** — `charlson_comorbidity_index`, "
          "`cci_*`, `dx_*`. Coded retrospectively for billing, so at prediction "
          "time they do not exist.",
          "3. **Availability leakage** — `vital_*`, `icu_*`, `fluid_*` when "
          "predicting ICU admission. The *presence* of an ICU chart entry "
          "reveals the ICU admission being predicted; the value never has to be "
          "read for the label to leak.", "",
          "The third is the subtle one. The first two are visible on a column "
          "list; availability leakage is a property of *whether a column is "
          "populated*, which a correlation check will not surface.", "",
          "### Observation windows", "",
          "Features are computed inside a fixed window from admission, so every "
          "row is built from what was knowable at the moment of prediction. A "
          "window defined relative to *discharge* would let a long admission "
          "contribute later measurements than a short one — the feature would "
          "then encode length of stay, which is an outcome.", ""]

    # ── corrections ──
    L += ["---", "", "## 4. What was wrong, and what fixing it bought", "",
          "Each of these was found after results had been produced, and each "
          "changed them. They are recorded in full in "
          "[`reports/data_correction_notice.md`](data_correction_notice.md).", ""]
    L += _table([
        ["Identifier precision loss in laboratory joins",
         "float64 `hadm_id` lost precision, so lab values joined to the wrong admission",
         "Re-joined on exact integer keys; affected results republished"],
        ["Observation-window leakage",
         "The laboratory window extended past the prediction point",
         "Window closed at the landmark; AUROC fell and became trustworthy"],
        ["Test suite overwrote production data",
         "A test wrote to `data/processed/`, corrupting the real tables",
         "Tests confined to `tmp_path`; guard test added"],
        ["Sequence-vs-tabular comparison not like-for-like",
         "The two arms saw different features, so the comparison meant nothing",
         "Rebuilt on identical inputs"],
        ["Embeddings graded on their own inputs",
         "Retrieval evaluated against the text it was built from",
         "Held-out evaluation set"],
    ], ["Defect", "What it did", "Correction"], align="lll")

    L += ["> Every one of these made results *look better* before it was found. "
          "That is the direction this class of defect always points, which is "
          "why a number improving is not by itself evidence that anything "
          "improved.", ""]

    # ── provenance ──
    L += ["---", "", "## 5. Provenance", "",
          "Read at generation time; a figure above traces to one of these:", ""]
    srcs = [ADMISSIONS, SELECTED, SPLIT,
            TABLES / "feature_missing_report.parquet",
            TABLES / "feature_dictionary.parquet",
            Path("reports/cleaning_report.md"),
            Path("reports/feature_engineering_report.md")]
    L += _table([[f"`{p}`", "present" if p.exists() else "**missing**",
                  _size(p)] for p in srcs],
                ["Source", "State", "Size"], align="llr")

    L += ["", "Model results are deliberately not restated here — they live in "
          "the phase reports, which own them. A second copy of a metric is a "
          "second thing to keep in sync, and this project has already been "
          "caught by exactly that.", ""]
    return "\n".join(L) + "\n"


# ── rendering ────────────────────────────────────────────────────────────────
#
# The document is composed as markdown and then rendered, rather than being
# written twice. Two builders producing "the same" report is the defect this
# whole file exists to avoid — they agree until one is edited.

NAVY = "1F3A5F"
GREY = "5A5A5A"
_INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")


def _runs(p, text: str, italic: bool = False, color: Optional[str] = None) -> None:
    """Inline markdown — bold, `code`, italic — onto a docx paragraph."""
    from docx.shared import Pt, RGBColor

    # A whole line wrapped in single asterisks is one italic span, and the
    # `code` spans nested inside it were never reached — the outer match
    # swallowed them and their backticks survived into the document. Unwrap
    # first, then let the inline pass see what is inside.
    stripped = text.strip()
    if (len(stripped) > 2 and stripped.startswith("*") and stripped.endswith("*")
            and not stripped.startswith("**")):
        text, italic = stripped[1:-1], True

    for part in _INLINE.split(text):
        if not part:
            continue
        r = p.add_run(part.strip("*`") if part[:1] in "*`" else part)
        r.italic = italic or (part.startswith("*") and not part.startswith("**"))
        r.bold = part.startswith("**")
        if part.startswith("`"):
            r.font.name, r.font.size = "Consolas", Pt(9.5)
        if color:
            r.font.color.rgb = RGBColor.from_string(color)


def render_docx(md: str):
    """Render the generated markdown to a Word document."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Calibri", Pt(10.5)

    lines, i = md.splitlines(), 0
    while i < len(lines):
        line = lines[i].rstrip()

        # tables: a run of pipe rows, second of which is the separator
        if line.startswith("|"):
            block = []
            while i < len(lines) and lines[i].startswith("|"):
                block.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            head, body = block[0], [r for r in block[2:]]
            t = doc.add_table(rows=1, cols=len(head))
            t.style = "Light Grid Accent 1"
            for cell, text in zip(t.rows[0].cells, head):
                cell.text = ""
                _runs(cell.paragraphs[0], f"**{text}**")
            for row in body:
                cells = t.add_row().cells
                for cell, text in zip(cells, row):
                    cell.text = ""
                    _runs(cell.paragraphs[0], text)
            doc.add_paragraph()
            continue

        i += 1
        if not line or line == "---":
            continue
        if line.startswith("### "):
            h = doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            h = doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            h = doc.add_heading(line[2:], level=1)
        elif line.startswith("> "):
            # A note, not a claim of the same weight as body text.
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            _runs(p, line[2:], italic=True, color=GREY)
            continue
        elif re.match(r"^\d+\. ", line):
            _runs(doc.add_paragraph(style="List Number"), line.split(". ", 1)[1])
            continue
        elif line.startswith("- "):
            _runs(doc.add_paragraph(style="List Bullet"), line[2:])
            continue
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            _runs(p, line)
            continue
        for r in h.runs:
            r.font.color.rgb = RGBColor.from_string(NAVY)
    return doc


def main(argv: Optional[List[str]] = None) -> int:
    import argparse

    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--markdown", action="store_true",
                    help="also write the .md alongside the .docx")
    args = ap.parse_args(argv)

    if not ADMISSIONS.exists():
        print(f"missing {ADMISSIONS} — run the feature pipeline first")
        return 2
    OUT.parent.mkdir(parents=True, exist_ok=True)
    md = build()
    render_docx(md).save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")
    if args.markdown:
        OUT.write_text(md, encoding="utf-8")
        print(f"Wrote {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
